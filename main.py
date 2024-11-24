import mysql.connector
import matplotlib
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from numpy.ma.extras import average
import scipy.stats as stats
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    passwd="Bobdole!",
    database="state_shoot_stats"
)
matplotlib.use('TkAgg')
figure = 1
table = 1
years = [2024, 2023, 2022, 2021, 2020]
friday_singles_events = [(2024, 6), (2023, 5), (2022, 5), (2021, 5), (2020, 6)]
friday_handicap_events = [(2024, 7), (2023, 6), (2022, 6), (2021, 6), (2020, 7)]
friday_doubles_events = [(2024, 5), (2023, 7), (2022, 7), (2021, 7), (2020, 5)]
championship_singles_events = [(2024, 8), (2023, 8), (2022, 8), (2021, 8), (2020, 8)]
championship_handicap_events = [(2024, 10), (2023, 10), (2022, 9), (2021, 9), (2020, 10)]
championship_doubles_events = [(2024, 9), (2023, 9), (2022, 7), (2021, 7), (2020,9)]

unique_totals = []
unique_residents = []
unique_non_residents = []
non_residents_by_state_labels = []
non_residents_by_state_values = []
youth_event_totals = []
mycursor = mydb.cursor()
resident_color = "lime"
non_resident_color = "red"
total_color = "blue"
best_fit_color = "magenta"


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def merge_row(table, row, columns):
    row = table.rows[row]
    cells_to_merge = row.cells[0:columns]
    merged_cell = cells_to_merge[0]
    merged_cell.merge(cells_to_merge[len(cells_to_merge) - 1])

def add_table(doc, data_frame, table_num, title, source):
    doc.add_paragraph()
    rows = len(data_frame.index) + 4
    columns = len(data_frame.columns)
    doc_table = doc.add_table(rows, columns)
    #table number
    merge_row(doc_table, 0, columns)
    doc_table.cell(0, 0).text = "Table {}".format(table_num)
    table_number_run = doc_table.cell(0,0).paragraphs[0].runs[0]
    table_number_run.font.name = "Times New Roman"
    table_number_run.font.size = Pt(12)
    table_number_run.font.bold = True
    #table_title
    merge_row(doc_table, 1, columns)
    doc_table.cell(1, 0).text = title.title()
    title_run = doc_table.cell(1, 0).paragraphs[0].runs[0]
    title_run.font.name = "Times New Roman"
    title_run.font.italic = True
    title_run.font.size = Pt(12)
    set_cell_border(doc_table.cell(1,0), bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"})
    #column headings
    headings = data_frame.columns.tolist()
    heading_counter = 0
    for heading in headings:
        cell = doc_table.cell(2, heading_counter)
        cell.text = heading
        paragraph = doc_table.cell(2, heading_counter).paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True
        heading_counter += 1
    #data
    for index, row in data_frame.iterrows():
        for column in range(columns):
            cell = doc_table.cell(index + 3, column)
            cell.text = str(row.iloc[column])
            paragraph = doc_table.cell(index + 3, column).paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    #source
    merge_row(doc_table, rows-1, columns)
    set_cell_border(doc_table.cell(rows-1, 0), top={"sz": 12, "val": "single", "color": "#000000", "space": "0"})
    source_cell = doc_table.cell(rows-1, 0)
    source_paragraph = source_cell.paragraphs[0]
    source_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    source_run_1 = source_paragraph.add_run()
    source_run_1.font.name = "Times New Roman"
    source_run_1.font.size = Pt(12)
    source_run_1.italic = True
    source_run_1.text = "Source: "
    source_run_2 = source_paragraph.add_run()
    source_run_2.font.name = "Times New Roman"
    source_run_2.font.size = Pt(12)
    source_run_2.italic = False
    source_run_2.text = source


# document = docx.Document()
# document.add_paragraph(str(years[0]) + "Summary Tables")
for year in years:
    mycursor.execute("select count(name) as name_count from (select distinct name from entries where year=%s) as names",
                     (year, ))
    unique_totals.append(mycursor.fetchall()[0][0])
    mycursor.execute("select count(name) as name_count from (select distinct name from entries where year=%s and state='ND') as names",
                     (year,))
    unique_residents.append(mycursor.fetchall()[0][0])
    mycursor.execute(
        "select count(name) as name_count from (select distinct name from entries where year=%s and state!='ND') as names",
        (year,))

    unique_non_residents.append(mycursor.fetchall()[0][0])
    mycursor.execute(
        "select distinct state, count(state) from (select distinct name, state from entries where year = %s and state != 'ND') as unique_shooters group by state order by state",
        (year,))
    temp_data = mycursor.fetchall()
    temp_states = []
    temp_counts = []
    for row in temp_data:
        temp_states.append(row[0])
        temp_counts.append(row[1])
    non_residents_by_state_labels.append(temp_states)
    non_residents_by_state_values.append(temp_counts)
mycursor.execute("select distinct year, count(year) from (select name, year from entries where event_number=1) as shooters group by year order by year desc")
youth_event_temp = mycursor.fetchall()



for row in youth_event_temp:
    youth_event_totals.append(row[1])
thursday_singles = []
mycursor.execute("select distinct year, count(year) from (select name, year from entries where event_number=2) as shooters group by year order by year desc")
thursday_singles_temp = mycursor.fetchall()
for row in thursday_singles_temp:
    thursday_singles.append(row[1])

thursday_handicap = []
mycursor.execute("select distinct year, count(year) from (select name, year from entries where event_number=3) as shooters group by year order by year desc")
thursday_handicap_temp = mycursor.fetchall()
for row in thursday_handicap_temp:
    thursday_handicap.append(row[1])

thursday_doubles = []
mycursor.execute("select distinct year, count(year) from (select name, year from entries where event_number=4) as shooters group by year order by year desc")
thursday_doubles_temp = mycursor.fetchall()
for row in thursday_doubles_temp:
    thursday_doubles.append(row[1])

friday_singles = []
first = True
friday_singles_query = ""
for row in friday_singles_events:
    if first:
        first = False
        friday_singles_query = "(select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])
    else:
        friday_singles_query += "union all (select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])

mycursor.execute(friday_singles_query)
friday_singles_temp = mycursor.fetchall()
for row in friday_singles_temp:
    friday_singles.append(row[2])
    
friday_handicap = []
first = True
friday_handicap_query = ""
for row in friday_handicap_events:
    if first:
        first = False
        friday_handicap_query = "(select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])
    else:
        friday_handicap_query += "union all (select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])

mycursor.execute(friday_handicap_query)
friday_handicap_temp = mycursor.fetchall()
for row in friday_handicap_temp:
    friday_handicap.append(row[2])
    
friday_doubles = []
first = True
friday_doubles_query = ""
for row in friday_doubles_events:
    if first:
        first = False
        friday_doubles_query = "(select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])
    else:
        friday_doubles_query += "union all (select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])

mycursor.execute(friday_doubles_query)
friday_doubles_temp = mycursor.fetchall()
for row in friday_doubles_temp:
    friday_doubles.append(row[2])
    
championship_singles = []
first = True
championship_singles_query = ""
for row in championship_singles_events:
    if first:
        first = False
        championship_singles_query = "(select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])
    else:
        championship_singles_query += "union all (select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])

mycursor.execute(championship_singles_query)
championship_singles_temp = mycursor.fetchall()
for row in championship_singles_temp:
    championship_singles.append(row[2])

championship_handicap = []
first = True
championship_handicap_query = ""
for row in championship_handicap_events:
    if first:
        first = False
        championship_handicap_query = "(select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])
    else:
        championship_handicap_query += "union all (select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])

mycursor.execute(championship_handicap_query)
championship_handicap_temp = mycursor.fetchall()
for row in championship_handicap_temp:
    championship_handicap.append(row[2])

championship_doubles = []
first = True
championship_doubles_query = ""
for row in championship_doubles_events:
    if first:
        first = False
        championship_doubles_query = "(select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])
    else:
        championship_doubles_query += "union all (select distinct year, event_number, count(year) from (select name, year, event_number from entries where event_number={0} and year={1}) as shooters group by year order by year desc)".format(row[1], row[0])

mycursor.execute(championship_doubles_query)
championship_doubles_temp = mycursor.fetchall()
for row in championship_doubles_temp:
    championship_doubles.append(row[2])

total_entries = []
mycursor.execute("select year, count(year) from entries group by year order by year desc")
total_entries_temp = mycursor.fetchall()
for row in total_entries_temp:
    total_entries.append(row[1])

total_unique_categories_years = []
total_unique_categories_labels = []
total_unique_categories_data =[]
mycursor.execute("select year, category, count(category) from (select distinct distinct name, year, category from entries order by name) as unique_shooters group by category, year order by year desc, category")
total_unique_temp = mycursor.fetchall()
for row in total_unique_temp:
    total_unique_categories_years.append(row[0])
    total_unique_categories_labels.append(row[1])
    total_unique_categories_data.append(row[2])
df_categories = pd.DataFrame({"Year": total_unique_categories_years, "Category": total_unique_categories_labels, "Counts": total_unique_categories_data})
df_categories_pivot= df_categories.pivot(index="Year", columns="Category", values="Counts")

total_unique_aggregate_categories_years = []
total_unique_aggregate_categories_labels = []
total_unique_aggregate_categories_data =[]
mycursor.execute("select year, aggregate_category, count(aggregate_category) from (select distinct distinct name, year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries order by name) as unique_shooters group by aggregate_category, year order by year desc, aggregate_category")
total_unique_aggregate_temp = mycursor.fetchall()
for row in total_unique_aggregate_temp:
    total_unique_aggregate_categories_years.append(row[0])
    total_unique_aggregate_categories_labels.append(row[1])
    total_unique_aggregate_categories_data.append(row[2])
df_categories_aggregate = pd.DataFrame({"Year": total_unique_aggregate_categories_years, "Category": total_unique_aggregate_categories_labels, "Counts": total_unique_aggregate_categories_data})
df_categories_aggregate_pivot= df_categories_aggregate.pivot(index="Year", columns="Category", values="Counts")



championship_singles_unique_category = []
championship_singles_unique_categories_years = []
championship_singles_unique_categories_labels = []
championship_singles_unique_categories_data =[]
first = True
championship_singles_unique_category_query = ""
for row in championship_singles_events:
    if first:
        first = False
        #(select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)
        championship_singles_unique_category_query = "(select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)".format(row[1], row[0])
    else:
        championship_singles_unique_category_query += " union all (select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)".format(row[1], row[0])
mycursor.execute(championship_singles_unique_category_query)
championship_singles_unique_category_temp = mycursor.fetchall()
for row in championship_singles_unique_category_temp:
    championship_singles_unique_categories_years.append(row[0])
    championship_singles_unique_categories_labels.append(row[1])
    championship_singles_unique_categories_data.append(row[2])
df_championship_singles_categories = pd.DataFrame({"Year": championship_singles_unique_categories_years, "Category": championship_singles_unique_categories_labels, "Counts": championship_singles_unique_categories_data})
df_championship_singles_categories_pivot= df_championship_singles_categories.pivot(index="Year", columns="Category", values="Counts")

championship_singles_aggregate_category = []
championship_singles_aggregate_categories_years = []
championship_singles_aggregate_categories_labels = []
championship_singles_aggregate_categories_data =[]
first = True
championship_singles_aggregate_category_query = ""
for row in championship_singles_events:
    if first:
        first = False
        championship_singles_aggregate_category_query = "(select year, aggregate_category, count(aggregate_category) from (select year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries where year={1} and event_number={0}) as unique_shooters group by aggregate_category order by aggregate_category)".format(row[1], row[0])
    else:
        championship_singles_aggregate_category_query += " union all (select year, aggregate_category, count(aggregate_category) from (select year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries where year={1} and event_number={0}) as unique_shooters group by aggregate_category order by aggregate_category)".format(row[1], row[0])
mycursor.execute(championship_singles_aggregate_category_query)
championship_singles_aggregate_category_temp = mycursor.fetchall()
for row in championship_singles_aggregate_category_temp:
    championship_singles_aggregate_categories_years.append(row[0])
    championship_singles_aggregate_categories_labels.append(row[1])
    championship_singles_aggregate_categories_data.append(row[2])
df_championship_singles_aggregate_categories = pd.DataFrame({"Year": championship_singles_aggregate_categories_years, "Category": championship_singles_aggregate_categories_labels, "Counts": championship_singles_aggregate_categories_data})
df_championship_singles_aggregate_categories_pivot= df_championship_singles_aggregate_categories.pivot(index="Year", columns="Category", values="Counts")

championship_handicap_unique_category = []
championship_handicap_unique_categories_years = []
championship_handicap_unique_categories_labels = []
championship_handicap_unique_categories_data =[]
first = True
aggregated_championship_handicap_unique_category_query = ""
for row in championship_handicap_events:
    if first:
        first = False
        aggregated_championship_handicap_unique_category_query = "(select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)".format(row[1], row[0])
    else:
        aggregated_championship_handicap_unique_category_query += " union all (select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)".format(row[1], row[0])
mycursor.execute(aggregated_championship_handicap_unique_category_query)
aggregated_championship_handicap_unique_category_temp = mycursor.fetchall()
for row in aggregated_championship_handicap_unique_category_temp:
    championship_handicap_unique_categories_years.append(row[0])
    championship_handicap_unique_categories_labels.append(row[1])
    championship_handicap_unique_categories_data.append(row[2])
df_championship_handicap_categories = pd.DataFrame({"Year": championship_handicap_unique_categories_years, "Category": championship_handicap_unique_categories_labels, "Counts": championship_handicap_unique_categories_data})
df_championship_handicap_categories_pivot= df_championship_handicap_categories.pivot(index="Year", columns="Category", values="Counts")

championship_handicap_aggregate_category = []
championship_handicap_aggregate_categories_years = []
championship_handicap_aggregate_categories_labels = []
championship_handicap_aggregate_categories_data =[]
first = True
championship_handicap_aggregate_category_query = ""
for row in championship_handicap_events:
    if first:
        first = False
        championship_handicap_aggregate_category_query = "(select year, aggregate_category, count(aggregate_category) from (select year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries where year={1} and event_number={0}) as unique_shooters group by aggregate_category order by aggregate_category)".format(row[1], row[0])
    else:
        championship_handicap_aggregate_category_query += " union all (select year, aggregate_category, count(aggregate_category) from (select year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries where year={1} and event_number={0}) as unique_shooters group by aggregate_category order by aggregate_category)".format(row[1], row[0])
mycursor.execute(championship_handicap_aggregate_category_query)
championship_handicap_aggregate_category_temp = mycursor.fetchall()
for row in championship_handicap_aggregate_category_temp:
    championship_handicap_aggregate_categories_years.append(row[0])
    championship_handicap_aggregate_categories_labels.append(row[1])
    championship_handicap_aggregate_categories_data.append(row[2])
df_championship_handicap_aggregate_categories = pd.DataFrame({"Year": championship_handicap_aggregate_categories_years, "Category": championship_handicap_aggregate_categories_labels, "Counts": championship_handicap_aggregate_categories_data})
df_championship_handicap_aggregate_categories_pivot= df_championship_handicap_aggregate_categories.pivot(index="Year", columns="Category", values="Counts")



championship_doubles_unique_category = []
championship_doubles_unique_categories_years = []
championship_doubles_unique_categories_labels = []
championship_doubles_unique_categories_data =[]
first = True
aggregated_championship_doubles_unique_category_query = ""
for row in championship_doubles_events:
    if first:
        first = False
        aggregated_championship_doubles_unique_category_query = "(select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)".format(row[1], row[0])
    else:
        aggregated_championship_doubles_unique_category_query += " union all (select year, category, count(category) from (select year, category from entries where year={1} and event_number={0}) as unique_shooters group by category order by category)".format(row[1], row[0])
mycursor.execute(aggregated_championship_doubles_unique_category_query)
aggregated_championship_doubles_unique_category_temp = mycursor.fetchall()
for row in aggregated_championship_doubles_unique_category_temp:
    championship_doubles_unique_categories_years.append(row[0])
    championship_doubles_unique_categories_labels.append(row[1])
    championship_doubles_unique_categories_data.append(row[2])
df_championship_doubles_categories = pd.DataFrame({"Year": championship_doubles_unique_categories_years, "Category": championship_doubles_unique_categories_labels, "Counts": championship_doubles_unique_categories_data})
df_championship_doubles_categories_pivot= df_championship_doubles_categories.pivot(index="Year", columns="Category", values="Counts")

championship_doubles_aggregate_category = []
championship_doubles_aggregate_categories_years = []
championship_doubles_aggregate_categories_labels = []
championship_doubles_aggregate_categories_data =[]
first = True
championship_doubles_aggregate_category_query = ""
for row in championship_doubles_events:
    if first:
        first = False
        championship_doubles_aggregate_category_query = "(select year, aggregate_category, count(aggregate_category) from (select year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries where year={1} and event_number={0}) as unique_shooters group by aggregate_category order by aggregate_category)".format(row[1], row[0])
    else:
        championship_doubles_aggregate_category_query += " union all (select year, aggregate_category, count(aggregate_category) from (select year, case when category = \"SJ\" or category=\"JR\" or category=\"JRG\" then \"Youth Category\" when category = \"LD1\" or category=\"LD1\" then \"Lady Category\" when category = \"SBVT\" or category=\"VT\" or category=\"SRVT\" then \"Veteran Category\"  else \"Open\" end as aggregate_category from entries where year={1} and event_number={0}) as unique_shooters group by aggregate_category order by aggregate_category)".format(row[1], row[0])
mycursor.execute(championship_doubles_aggregate_category_query)
championship_doubles_aggregate_category_temp = mycursor.fetchall()
for row in championship_doubles_aggregate_category_temp:
    championship_doubles_aggregate_categories_years.append(row[0])
    championship_doubles_aggregate_categories_labels.append(row[1])
    championship_doubles_aggregate_categories_data.append(row[2])
df_championship_doubles_aggregate_categories = pd.DataFrame({"Year": championship_doubles_aggregate_categories_years, "Category": championship_doubles_aggregate_categories_labels, "Counts": championship_doubles_aggregate_categories_data})
df_championship_doubles_aggregate_categories_pivot= df_championship_doubles_aggregate_categories.pivot(index="Year", columns="Category", values="Counts")

youth_event_categories_labels = []
youth_event_categories_counts = []
mycursor.execute("select category, count(category) from entries where year = 2024 and event_number = 1 group by category order by category")
youth_event_categories_temp = mycursor.fetchall()
for row in youth_event_categories_temp:
    youth_event_categories_labels.append(row[0])
    youth_event_categories_counts.append(row[1])

youth_event_classes_labels = []
youth_event_classes_counts = []
mycursor.execute("select class, count(class) from entries where year = 2024 and event_number = 1 group by class order by class")
youth_event_classes_temp = mycursor.fetchall()
for row in youth_event_classes_temp:
    youth_event_classes_labels.append(row[0])
    youth_event_classes_counts.append(row[1])

youth_event_categories_hist_scores = []
for category in youth_event_categories_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 1 and category = \"{}\" order by category".format(category))
    youth_event_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in youth_event_categories_hist_temp:
        temp_list.append(row[0])
    youth_event_categories_hist_scores.append(temp_list)

youth_event_classes_hist_scores = []
for class_ in youth_event_classes_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 1 and class = \"{}\" order by class".format(class_))
    youth_event_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in youth_event_classes_hist_temp:
        temp_list.append(row[0])
    youth_event_classes_hist_scores.append(temp_list)

thursday_singles_categories_labels = []
thursday_singles_categories_counts = []
mycursor.execute("select category, count(category) from entries where year = 2024 and event_number = 2 group by category order by category")
thursday_singles_categories_temp = mycursor.fetchall()
for row in thursday_singles_categories_temp:
    thursday_singles_categories_labels.append(row[0])
    thursday_singles_categories_counts.append(row[1])

thursday_singles_classes_labels = []
thursday_singles_classes_counts = []
mycursor.execute("select class, count(class) from entries where year = 2024 and event_number = 2 group by class order by class")
thursday_singles_classes_temp = mycursor.fetchall()
for row in thursday_singles_classes_temp:
    thursday_singles_classes_labels.append(row[0])
    thursday_singles_classes_counts.append(row[1])

thursday_singles_categories_hist_scores = []
for category in thursday_singles_categories_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 2 and category = \"{}\" order by category".format(category))
    thursday_singles_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in thursday_singles_categories_hist_temp:
        temp_list.append(row[0])
    thursday_singles_categories_hist_scores.append(temp_list)

thursday_singles_classes_hist_scores = []
for class_ in thursday_singles_classes_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 2 and class = \"{}\" order by class".format(class_))
    thursday_singles_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in thursday_singles_classes_hist_temp:
        temp_list.append(row[0])
    thursday_singles_classes_hist_scores.append(temp_list)

thursday_handicap_categories_labels = []
thursday_handicap_categories_counts = []
mycursor.execute("select category, count(category) from entries where year = 2024 and event_number = 3 group by category order by category")
thursday_handicap_categories_temp = mycursor.fetchall()
for row in thursday_handicap_categories_temp:
    thursday_handicap_categories_labels.append(row[0])
    thursday_handicap_categories_counts.append(row[1])

thursday_handicap_yardages_labels = []
thursday_handicap_yardages_counts = []
mycursor.execute("select yardage, count(yardage) from entries where year = 2024 and event_number = 3 group by yardage order by yardage")
thursday_handicap_yardages_temp = mycursor.fetchall()
for row in thursday_handicap_yardages_temp:
    thursday_handicap_yardages_labels.append(row[0])
    thursday_handicap_yardages_counts.append(row[1])

thursday_handicap_categories_hist_scores = []
for category in thursday_handicap_categories_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 3 and category = \"{}\" order by category".format(category))
    thursday_handicap_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in thursday_handicap_categories_hist_temp:
        temp_list.append(row[0])
    thursday_handicap_categories_hist_scores.append(temp_list)

thursday_handicap_yardages_hist_scores = []
for class_ in thursday_handicap_yardages_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 3 and yardage = \"{}\" order by yardage".format(class_))
    thursday_handicap_yardages_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in thursday_handicap_yardages_hist_temp:
        temp_list.append(row[0])
    thursday_handicap_yardages_hist_scores.append(temp_list)

thursday_doubles_categories_labels = []
thursday_doubles_categories_counts = []
mycursor.execute("select category, count(category) from entries where year = 2024 and event_number = 4 group by category order by category")
thursday_doubles_categories_temp = mycursor.fetchall()
for row in thursday_doubles_categories_temp:
    thursday_doubles_categories_labels.append(row[0])
    thursday_doubles_categories_counts.append(row[1])
    
thursday_doubles_classes_labels = []
thursday_doubles_classes_counts = []
mycursor.execute("select class, count(class) from entries where year = 2024 and event_number = 4 group by class order by class")
thursday_doubles_classes_temp = mycursor.fetchall()
for row in thursday_doubles_classes_temp:
    thursday_doubles_classes_labels.append(row[0])
    thursday_doubles_classes_counts.append(row[1])

thursday_doubles_categories_hist_scores = []
for category in thursday_doubles_categories_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 4 and category = \"{}\" order by category".format(category))
    thursday_doubles_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in thursday_doubles_categories_hist_temp:
        temp_list.append(row[0])
    thursday_doubles_categories_hist_scores.append(temp_list)

thursday_doubles_classes_hist_scores = []
for class_ in thursday_doubles_classes_labels:
    mycursor.execute("select total_score from entries where year = 2024 and event_number = 4 and class = \"{}\" order by class".format(class_))
    thursday_doubles_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in thursday_doubles_classes_hist_temp:
        temp_list.append(row[0])
    thursday_doubles_classes_hist_scores.append(temp_list)

friday_doubles_categories_labels = []
friday_doubles_categories_counts = []
mycursor.execute(
    "select category, count(category) from entries where year = 2024 and event_number = 5 group by category order by category")
friday_doubles_categories_temp = mycursor.fetchall()
for row in friday_doubles_categories_temp:
    friday_doubles_categories_labels.append(row[0])
    friday_doubles_categories_counts.append(row[1])

friday_doubles_classes_labels = []
friday_doubles_classes_counts = []
mycursor.execute(
    "select class, count(class) from entries where year = 2024 and event_number = 5 group by class order by class")
friday_doubles_classes_temp = mycursor.fetchall()
for row in friday_doubles_classes_temp:
    friday_doubles_classes_labels.append(row[0])
    friday_doubles_classes_counts.append(row[1])

friday_doubles_categories_hist_scores = []
for category in friday_doubles_categories_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 5 and category = \"{}\" order by category".format(
            category))
    friday_doubles_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in friday_doubles_categories_hist_temp:
        temp_list.append(row[0])
    friday_doubles_categories_hist_scores.append(temp_list)

friday_doubles_classes_hist_scores = []
for class_ in friday_doubles_classes_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 5 and class = \"{}\" order by class".format(
            class_))
    friday_doubles_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in friday_doubles_classes_hist_temp:
        temp_list.append(row[0])
    friday_doubles_classes_hist_scores.append(temp_list)

friday_singles_categories_labels = []
friday_singles_categories_counts = []
mycursor.execute(
    "select category, count(category) from entries where year = 2024 and event_number = 6 group by category order by category")
friday_singles_categories_temp = mycursor.fetchall()
for row in friday_singles_categories_temp:
    friday_singles_categories_labels.append(row[0])
    friday_singles_categories_counts.append(row[1])

friday_singles_classes_labels = []
friday_singles_classes_counts = []
mycursor.execute(
    "select class, count(class) from entries where year = 2024 and event_number = 6 group by class order by class")
friday_singles_classes_temp = mycursor.fetchall()
for row in friday_singles_classes_temp:
    friday_singles_classes_labels.append(row[0])
    friday_singles_classes_counts.append(row[1])

friday_singles_categories_hist_scores = []
for category in friday_singles_categories_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 6 and category = \"{}\" order by category".format(
            category))
    friday_singles_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in friday_singles_categories_hist_temp:
        temp_list.append(row[0])
    friday_singles_categories_hist_scores.append(temp_list)

friday_singles_classes_hist_scores = []
for class_ in friday_singles_classes_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 6 and class = \"{}\" order by class".format(
            class_))
    friday_singles_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in friday_singles_classes_hist_temp:
        temp_list.append(row[0])
    friday_singles_classes_hist_scores.append(temp_list)

friday_handicap_categories_labels = []
friday_handicap_categories_counts = []
mycursor.execute(
    "select category, count(category) from entries where year = 2024 and event_number = 7 group by category order by category")
friday_handicap_categories_temp = mycursor.fetchall()
for row in friday_handicap_categories_temp:
    friday_handicap_categories_labels.append(row[0])
    friday_handicap_categories_counts.append(row[1])

friday_handicap_yardages_labels = []
friday_handicap_yardages_counts = []
mycursor.execute(
    "select yardage, count(yardage) from entries where year = 2024 and event_number = 7 group by yardage order by yardage")
friday_handicap_yardages_temp = mycursor.fetchall()
for row in friday_handicap_yardages_temp:
    friday_handicap_yardages_labels.append(row[0])
    friday_handicap_yardages_counts.append(row[1])

friday_handicap_categories_hist_scores = []
for category in friday_handicap_categories_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 7 and category = \"{}\" order by category".format(
            category))
    friday_handicap_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in friday_handicap_categories_hist_temp:
        temp_list.append(row[0])
    friday_handicap_categories_hist_scores.append(temp_list)

friday_handicap_yardages_hist_scores = []
for yardage_ in friday_handicap_yardages_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 7 and yardage = \"{}\" order by yardage".format(
            yardage_))
    friday_handicap_yardages_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in friday_handicap_yardages_hist_temp:
        temp_list.append(row[0])
    friday_handicap_yardages_hist_scores.append(temp_list)
    
championship_singles_categories_labels = []
championship_singles_categories_counts = []
mycursor.execute(
    "select category, count(category) from entries where year = 2024 and event_number = 8 group by category order by category")
championship_singles_categories_temp = mycursor.fetchall()
for row in championship_singles_categories_temp:
    championship_singles_categories_labels.append(row[0])
    championship_singles_categories_counts.append(row[1])

championship_singles_classes_labels = []
championship_singles_classes_counts = []
mycursor.execute(
    "select class, count(class) from entries where year = 2024 and event_number = 8 group by class order by class")
championship_singles_classes_temp = mycursor.fetchall()
for row in championship_singles_classes_temp:
    championship_singles_classes_labels.append(row[0])
    championship_singles_classes_counts.append(row[1])

championship_singles_categories_hist_scores = []
for category in championship_singles_categories_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 8 and category = \"{}\" order by category".format(
            category))
    championship_singles_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in championship_singles_categories_hist_temp:
        temp_list.append(row[0])
    championship_singles_categories_hist_scores.append(temp_list)

championship_singles_classes_hist_scores = []
for class_ in championship_singles_classes_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 8 and class = \"{}\" order by class".format(
            class_))
    championship_singles_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in championship_singles_classes_hist_temp:
        temp_list.append(row[0])
    championship_singles_classes_hist_scores.append(temp_list)

championship_doubles_categories_labels = []
championship_doubles_categories_counts = []
mycursor.execute(
    "select category, count(category) from entries where year = 2024 and event_number = 9 group by category order by category")
championship_doubles_categories_temp = mycursor.fetchall()
for row in championship_doubles_categories_temp:
    championship_doubles_categories_labels.append(row[0])
    championship_doubles_categories_counts.append(row[1])

championship_doubles_classes_labels = []
championship_doubles_classes_counts = []
mycursor.execute(
    "select class, count(class) from entries where year = 2024 and event_number = 9 group by class order by class")
championship_doubles_classes_temp = mycursor.fetchall()
for row in championship_doubles_classes_temp:
    championship_doubles_classes_labels.append(row[0])
    championship_doubles_classes_counts.append(row[1])

championship_doubles_categories_hist_scores = []
for category in championship_doubles_categories_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 9 and category = \"{}\" order by category".format(
            category))
    championship_doubles_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in championship_doubles_categories_hist_temp:
        temp_list.append(row[0])
    championship_doubles_categories_hist_scores.append(temp_list)

championship_doubles_classes_hist_scores = []
for class_ in championship_doubles_classes_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 9 and class = \"{}\" order by class".format(
            class_))
    championship_doubles_classes_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in championship_doubles_classes_hist_temp:
        temp_list.append(row[0])
    championship_doubles_classes_hist_scores.append(temp_list)

championship_handicap_categories_labels = []
championship_handicap_categories_counts = []
mycursor.execute(
    "select category, count(category) from entries where year = 2024 and event_number = 10 group by category order by category")
championship_handicap_categories_temp = mycursor.fetchall()
for row in championship_handicap_categories_temp:
    championship_handicap_categories_labels.append(row[0])
    championship_handicap_categories_counts.append(row[1])

championship_handicap_yardages_labels = []
championship_handicap_yardages_counts = []
mycursor.execute(
    "select yardage, count(yardage) from entries where year = 2024 and event_number = 10 group by yardage order by yardage")
championship_handicap_yardages_temp = mycursor.fetchall()
for row in championship_handicap_yardages_temp:
    championship_handicap_yardages_labels.append(row[0])
    championship_handicap_yardages_counts.append(row[1])

championship_handicap_categories_hist_scores = []
for category in championship_handicap_categories_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 10 and category = \"{}\" order by category".format(
            category))
    championship_handicap_categories_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in championship_handicap_categories_hist_temp:
        temp_list.append(row[0])
    championship_handicap_categories_hist_scores.append(temp_list)

championship_handicap_yardages_hist_scores = []
for yardage_ in championship_handicap_yardages_labels:
    mycursor.execute(
        "select total_score from entries where year = 2024 and event_number = 10 and yardage = \"{}\" order by yardage".format(
            yardage_))
    championship_handicap_yardages_hist_temp = mycursor.fetchall()
    temp_list = []
    for row in championship_handicap_yardages_hist_temp:
        temp_list.append(row[0])
    championship_handicap_yardages_hist_scores.append(temp_list)

mycursor.execute("select year, singles, handicap, doubles, total from championship_entries where state='North Dakota' order by year")
temp_data = mycursor.fetchall()
temp_year = []
temp_singles = []
temp_handicap = []
temp_doubles = []
temp_total = []
for row in temp_data:
    temp_year.append(row[0])
    temp_singles.append(row[1])
    temp_handicap.append(row[2])
    temp_doubles.append(row[3])
    temp_total.append(row[4])
df_extended_years_nd =pd.DataFrame({"Year": temp_year, "Singles": temp_singles, "Handicap": temp_handicap, "Doubles": temp_doubles, "Total": temp_total})

temp_years = []
temp_states = []
temp_entries =[]
mycursor.execute("select year, state, singles from championship_entries where state='North Dakota' or state='Minnesota' or state='South Dakota' or state='Montana' or state='Manitoba' or state='Saskatchewan' order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_states.append(row[1])
    temp_entries.append(row[2])
df_extended_years_adjacent_singles = pd.DataFrame({"Year": temp_years, "State": temp_states, "Entries": temp_entries})
df_extended_years_adjacent_singles_pivot= df_extended_years_adjacent_singles.pivot(index="Year", columns="State", values="Entries")
temp_averages = df_extended_years_adjacent_singles_pivot.mean(axis=1)
df_extended_years_adjacent_singles_pivot["Average"] = temp_averages

temp_years = []
temp_states = []
temp_entries =[]
mycursor.execute("select year, state, handicap from championship_entries where state='North Dakota' or state='Minnesota' or state='South Dakota' or state='Montana' or state='Manitoba' or state='Saskatchewan' order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_states.append(row[1])
    temp_entries.append(row[2])
df_extended_years_adjacent_handicap = pd.DataFrame({"Year": temp_years, "State": temp_states, "Entries": temp_entries})
df_extended_years_adjacent_handicap_pivot= df_extended_years_adjacent_handicap.pivot(index="Year", columns="State", values="Entries")
temp_averages = df_extended_years_adjacent_handicap_pivot.mean(axis=1)
df_extended_years_adjacent_handicap_pivot["Average"] = temp_averages

temp_years = []
temp_states = []
temp_entries =[]
mycursor.execute("select year, state, doubles from championship_entries where state='North Dakota' or state='Minnesota' or state='South Dakota' or state='Montana' or state='Manitoba' or state='Saskatchewan' order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_states.append(row[1])
    temp_entries.append(row[2])
df_extended_years_adjacent_doubles = pd.DataFrame({"Year": temp_years, "State": temp_states, "Entries": temp_entries})
df_extended_years_adjacent_doubles_pivot= df_extended_years_adjacent_doubles.pivot(index="Year", columns="State", values="Entries")
temp_averages = df_extended_years_adjacent_doubles_pivot.mean(axis=1)
df_extended_years_adjacent_doubles_pivot["Average"] = temp_averages

temp_years = []
temp_states = []
temp_entries =[]
mycursor.execute("select year, state, total from championship_entries where state='North Dakota' or state='Minnesota' or state='South Dakota' or state='Montana' or state='Manitoba' or state='Saskatchewan' order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_states.append(row[1])
    temp_entries.append(row[2])
df_extended_years_adjacent_total = pd.DataFrame({"Year": temp_years, "State": temp_states, "Entries": temp_entries})
df_extended_years_adjacent_total_pivot= df_extended_years_adjacent_total.pivot(index="Year", columns="State", values="Entries")
temp_averages = df_extended_years_adjacent_total_pivot.mean(axis=1)
df_extended_years_adjacent_total_pivot["Average"] = temp_averages

#todo: remove restriction on 2024 when ATA average book comes out
temp_years = []
temp_zones = []
temp_entries =[]
mycursor.execute("select year, zone, avg(singles) from (select year, zones.zone, singles from championship_entries INNER JOIN zones on championship_entries.state = zones.state where year != 2024) as zone_mask group by zone, year order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_zones.append(row[1])
    temp_entries.append(row[2])

df_extended_years_all_singles = pd.DataFrame({"Year": temp_years, "Zone": temp_zones, "Entries": temp_entries})
df_extended_years_all_singles_pivot= df_extended_years_all_singles.pivot(index="Year", columns="Zone", values="Entries")
df_extended_years_all_singles_years = temp_years
mycursor.execute("select year, avg(singles) from championship_entries where year != 2024 group by year order by year")
temp_data = mycursor.fetchall()
temp_average = []
temp_years = []
for row in temp_data:
    temp_years.append(row[0])
    temp_average.append(row[1])
df_extended_years_all_singles_pivot['Average'] = temp_average
extended_years_all_singles_pivot_average_years = temp_years
temp_nd = []
mycursor.execute("select year, singles from championship_entries where state = 'North Dakota' and year !=2024 order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_nd.append(row[1])
df_extended_years_all_singles_pivot['North Dakota'] = temp_nd

#todo: remove restriction on 2024 when ATA average book comes out
temp_years = []
temp_zones = []
temp_entries =[]
mycursor.execute("select year, zone, avg(handicap) from (select year, zones.zone, handicap from championship_entries INNER JOIN zones on championship_entries.state = zones.state where year != 2024) as zone_mask group by zone, year order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_zones.append(row[1])
    temp_entries.append(row[2])

df_extended_years_all_handicap = pd.DataFrame({"Year": temp_years, "Zone": temp_zones, "Entries": temp_entries})
df_extended_years_all_handicap_pivot= df_extended_years_all_handicap.pivot(index="Year", columns="Zone", values="Entries")
df_extended_years_all_handicap_years = temp_years
mycursor.execute("select year, avg(handicap) from championship_entries where year != 2024 group by year order by year")
temp_data = mycursor.fetchall()
temp_average = []
temp_years = []
for row in temp_data:
    temp_years.append(row[0])
    temp_average.append(row[1])
df_extended_years_all_handicap_pivot['Average'] = temp_average
extended_years_all_handicap_pivot_average_years = temp_years
temp_nd = []
mycursor.execute("select year, handicap from championship_entries where state = 'North Dakota' and year !=2024 order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_nd.append(row[1])
df_extended_years_all_handicap_pivot['North Dakota'] = temp_nd

#todo: remove restriction on 2024 when ATA average book comes out
temp_years = []
temp_zones = []
temp_entries =[]
mycursor.execute("select year, zone, avg(doubles) from (select year, zones.zone, doubles from championship_entries INNER JOIN zones on championship_entries.state = zones.state where year != 2024) as zone_mask group by zone, year order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_zones.append(row[1])
    temp_entries.append(row[2])

df_extended_years_all_doubles = pd.DataFrame({"Year": temp_years, "Zone": temp_zones, "Entries": temp_entries})
df_extended_years_all_doubles_pivot= df_extended_years_all_doubles.pivot(index="Year", columns="Zone", values="Entries")
df_extended_years_all_doubles_years = temp_years
mycursor.execute("select year, avg(doubles) from championship_entries where year != 2024 group by year order by year")
temp_data = mycursor.fetchall()
temp_average = []
temp_years = []
for row in temp_data:
    temp_years.append(row[0])
    temp_average.append(row[1])
df_extended_years_all_doubles_pivot['Average'] = temp_average
extended_years_all_doubles_pivot_average_years = temp_years
temp_nd = []
mycursor.execute("select year, doubles from championship_entries where state = 'North Dakota' and year !=2024 order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_nd.append(row[1])
df_extended_years_all_doubles_pivot['North Dakota'] = temp_nd

#todo: remove restriction on 2024 when ATA average book comes out
temp_years = []
temp_zones = []
temp_entries =[]
mycursor.execute("select year, zone, avg(total) from (select year, zones.zone, total from championship_entries INNER JOIN zones on championship_entries.state = zones.state where year != 2024) as zone_mask group by zone, year order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_years.append(row[0])
    temp_zones.append(row[1])
    temp_entries.append(row[2])

df_extended_years_all_total = pd.DataFrame({"Year": temp_years, "Zone": temp_zones, "Entries": temp_entries})
df_extended_years_all_total_pivot= df_extended_years_all_total.pivot(index="Year", columns="Zone", values="Entries")
df_extended_years_all_total_years = temp_years
mycursor.execute("select year, avg(total) from championship_entries where year != 2024 group by year order by year")
temp_data = mycursor.fetchall()
temp_average = []
temp_years = []
for row in temp_data:
    temp_years.append(row[0])
    temp_average.append(row[1])
df_extended_years_all_total_pivot['Average'] = temp_average
extended_years_all_total_pivot_average_years = temp_years
temp_nd = []
mycursor.execute("select year, total from championship_entries where state = 'North Dakota' and year !=2024 order by year")
temp_data = mycursor.fetchall()
for row in temp_data:
    temp_nd.append(row[1])
df_extended_years_all_total_pivot['North Dakota'] = temp_nd

plt.figure()
#Figure 1: plot without best fit lines
plt.plot(years, unique_totals, label="Total Shooters", marker='o', color=total_color)
plt.plot(years, unique_residents, label="Residents", marker='o', color=resident_color)
plt.plot(years, unique_non_residents, label="Non-Residents", marker='o', color=non_resident_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Unique Shooters by Residency".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#figure 2: best fit lines added to Figure 1.
plt.figure()
plt.plot(years, unique_totals, label="Total Shooters", marker='o', color=total_color)
m_t, b_t = np.polyfit(years, unique_totals, 1)
x = np.array(years)
y_t = m_t * x + b_t
plt.plot(x, y_t, color=best_fit_color)
plt.text(average(years), min(unique_totals)-10, "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_t,b_t, 1))
plt.plot(years, unique_residents, label="Residents", marker='o', color=resident_color)
m_r, b_r = np.polyfit(years, unique_residents, 1)
y_r = m_r * x + b_r
plt.plot(x, y_r, color=best_fit_color)
plt.text(average(years), min(unique_residents)-10, "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_r,b_r, 1))
plt.plot(years, unique_non_residents, label="Non-Residents", marker='o', color=non_resident_color)
m_n, b_n = np.polyfit(years, unique_non_residents, 1)
y_n = m_n * x + b_n
plt.plot(x, y_n, color=best_fit_color)
plt.text(average(years), min(unique_non_residents)-10, "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_n,b_n, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Unique Shooters by Residency With Best Fit Lines".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 3: pie charts with residents vs. non-residents
for i in range(0, len(years)):
    plt.figure()
    labels = ["Residents", "Non-Residents"]
    colors = [resident_color, non_resident_color]
    data = [unique_residents[i], unique_non_residents[i]]
    plt.pie(data, labels=labels, colors=colors, autopct="%1.1f%%")
    plt.title("Figure {}: {} Unique Attendance by Residency".format(figure, years[i]))
    plt.savefig("figure{}.png".format(figure))
    figure +=1
plt.close('all')

#Figure 4: pie charts with non-residents by State
for i in range(0, len(years)):
    plt.figure()
    plt.pie(np.array(non_residents_by_state_values[i]), labels=np.array(non_residents_by_state_labels[i]), autopct="%1.1f%%")
    plt.title("Figure {}: {} Unique Non-Resident Attendance by State".format(figure, years[i]))
    plt.savefig("figure{}.png".format(figure))
    figure +=1
plt.close('all')

#Figure 5: Youth Event plot without best fit lines
plt.figure()
plt.plot(years, youth_event_totals, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Youth Event Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 6: Youth Event plot with best fit lines
plt.figure()
plt.plot(years, youth_event_totals, label="Total Shooters", marker='o', color=total_color)
m_youth, b_youth = np.polyfit(years, youth_event_totals, 1)
y_youth = m_youth * x + b_youth
plt.plot(x, y_youth, color=best_fit_color)
plt.text(average(years), min(youth_event_totals), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_youth,b_youth, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Youth Event Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 7: Thursday Singles plot without best fit lines
plt.figure()
plt.plot(years, thursday_singles, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Thursday Singles Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 8: Thursday Singles plot with best fit lines
plt.figure()
plt.plot(years, thursday_singles, label="Total Shooters", marker='o', color=total_color)
m_ts, b_ts = np.polyfit(years, thursday_singles, 1)
y_ts = m_ts * x + b_ts
plt.plot(x, y_ts, color=best_fit_color)
plt.text(average(years), min(thursday_singles), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_ts,b_ts, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Thursday Singles Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 9: Thursday Handicap plot without best fit lines
plt.figure()
plt.plot(years, thursday_handicap, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Thursday Handicap Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 10: Thursday handicap plot with best fit lines
plt.figure()
plt.plot(years, thursday_handicap, label="Total Shooters", marker='o', color=total_color)
m_th, b_th = np.polyfit(years, thursday_handicap, 1)
y_th = m_th * x + b_th
plt.plot(x, y_th, color=best_fit_color)
plt.text(average(years), min(thursday_handicap), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_th,b_th, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Thursday Handicap Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 11: Thursday Doubles plot without best fit lines
plt.figure()
plt.plot(years, thursday_doubles, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Thursday Doubles Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 12: Thursday doubles plot with best fit lines
plt.figure()
plt.plot(years, thursday_doubles, label="Total Shooters", marker='o', color=total_color)
m_th, b_th = np.polyfit(years, thursday_doubles, 1)
y_th = m_th * x + b_th
plt.plot(x, y_th, color=best_fit_color)
plt.text(average(years), min(thursday_doubles), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m_th,b_th, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Thursday Doubles Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 13: Friday Singles plot without best fit lines
plt.figure()
plt.plot(years, friday_singles, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Friday Singles Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 14: Friday singles plot with best fit lines
plt.figure()
plt.plot(years, friday_singles, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, friday_singles, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(friday_singles), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Friday Singles Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 15: Friday handicap plot without best fit lines
plt.figure()
plt.plot(years, friday_handicap, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Friday Handicap Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 16: Friday handicap plot with best fit lines
plt.figure()
plt.plot(years, friday_handicap, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, friday_handicap, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(friday_handicap), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Friday Handicap Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 17: Friday doubles plot without best fit lines
plt.figure()
plt.plot(years, friday_doubles, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Friday Doubles Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 18: Friday doubles plot with best fit lines
plt.figure()
plt.plot(years, friday_doubles, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, friday_doubles, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(friday_doubles), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Friday Doubles Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 19: Championship singles plot without best fit lines
plt.figure()
plt.plot(years, championship_singles, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Championship Singles Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 20: Championship singles plot with best fit lines
plt.figure()
plt.plot(years, championship_singles, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, championship_singles, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(championship_singles), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total category_championship Singles Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 21: category_championship handicap plot without best fit lines
plt.figure()
plt.plot(years, championship_handicap, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total category_championship Handicap Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 22: Championship handicap plot with best fit lines
plt.figure()
plt.plot(years, championship_handicap, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, championship_handicap, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(championship_handicap), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Championship Handicap Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 23: Championship doubles plot without best fit lines
plt.figure()
plt.plot(years, championship_doubles, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Championship Doubles Shooters".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 24: category_championship doubles plot with best fit lines
plt.figure()
plt.plot(years, championship_doubles, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, championship_doubles, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(championship_doubles), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Championship Doubles Shooters With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 25: Total Entries plot without best fit lines
plt.figure()
plt.plot(years, total_entries, label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Event Entries".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 26: Total Entries plot with best fit lines
plt.figure()
plt.plot(years, total_entries, label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(years, total_entries, 1)
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(years), min(total_entries), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Entries With Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 27: Total Categories plot with best fit lines
plt.figure()
df_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Unique Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 28: Aggregated Categories plot with best fit lines
plt.figure()
df_categories_aggregate_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Total Unique Aggregated Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 29: Championship Singles Categories plot
plt.figure()
df_championship_singles_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Championship Singles Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 29.1: Championship Singles Aggregated Categories plot
plt.figure()
df_championship_singles_aggregate_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Championship Singles Aggregated Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 30: Championship Handicap Categories plot
plt.figure()
df_championship_handicap_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Championship Handicap Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 30.1: Championship handicap Aggregated Categories plot
plt.figure()
df_championship_handicap_aggregate_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Championship Handicap Aggregated Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 31: Championship Doubles Categories plot
plt.figure()
df_championship_doubles_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Championship Doubles Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 31.1: Championship doubles Aggregated Categories plot
plt.figure()
df_championship_doubles_aggregate_categories_pivot.plot()
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(years)
plt.title("Figure {}: Championship Doubles Aggregated Category Shooters By Year".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 32 Youth Event Categories
plt.figure()
plt.pie(youth_event_categories_counts, labels=youth_event_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Youth Event (Event 1) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 33 Youth Event Classes
plt.figure()
plt.pie(youth_event_classes_counts, labels=youth_event_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Youth Event (Event 1) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 33 Youth Event Category Histogram
plt.figure()
fig, axs = plt.subplots(len(youth_event_categories_labels), 1, sharex=True)
for i in range(0,len(youth_event_categories_labels)):
    axs[i].hist(youth_event_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(youth_event_categories_labels[i], np.mean(youth_event_categories_hist_scores[i]), np.std(youth_event_categories_hist_scores[i])))
    if youth_event_categories_labels[i] != "JRG":
        data = np.array(youth_event_categories_hist_scores[i])
        a, loc, scale = stats.skewnorm.fit(data)
        x = np.linspace(0, 100)
        pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
        axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Youth Event (Event 1) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 34 Youth Event Class Histogram
plt.figure()
fig, axs = plt.subplots(len(youth_event_classes_labels), 1, sharex=True)
for i in range(0,len(youth_event_classes_labels)):
    axs[i].hist(youth_event_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(youth_event_classes_labels[i], np.mean(youth_event_classes_hist_scores[i]), np.std(youth_event_classes_hist_scores[i])))
    data = np.array(youth_event_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Youth Event (Event 1) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 35 Thursday Singles Categories
plt.figure()
plt.pie(thursday_singles_categories_counts, labels=thursday_singles_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Thursday Singles (Event 2) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 36 Thursday Singles Classes
plt.figure()
plt.pie(thursday_singles_classes_counts, labels=thursday_singles_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Thursday Singles (Event 2) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 37 Thursday Singles Category Histogram
plt.figure()
fig, axs = plt.subplots(len(thursday_singles_categories_labels), 1, sharex=True)
for i in range(0,len(thursday_singles_categories_labels)):
    axs[i].hist(thursday_singles_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(thursday_singles_categories_labels[i], np.mean(thursday_singles_categories_hist_scores[i]), np.std(thursday_singles_categories_hist_scores[i])))
    data = np.array(thursday_singles_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Thursday Singles (Event 2) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 38 Thursday Singles Class Histogram
plt.figure()
fig, axs = plt.subplots(len(thursday_singles_classes_labels), 1, sharex=True)
for i in range(0,len(thursday_singles_classes_labels)):
    axs[i].hist(thursday_singles_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(thursday_singles_classes_labels[i], np.mean(thursday_singles_classes_hist_scores[i]), np.std(thursday_singles_classes_hist_scores[i])))
    data = np.array(thursday_singles_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Thursday Singles (Event 2) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 39 Thursday handicap Categories
plt.figure()
plt.pie(thursday_handicap_categories_counts, labels=thursday_handicap_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Thursday Handicap (Event 3) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 40 Thursday handicap yardages
plt.figure()
plt.pie(thursday_handicap_yardages_counts, labels=thursday_handicap_yardages_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Thursday Handicap (Event 3) Shooters by Yardage".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 41 Thursday handicap Category Histogram
plt.figure()
fig, axs = plt.subplots(len(thursday_handicap_categories_labels), 1, sharex=True)
for i in range(0,len(thursday_handicap_categories_labels)):
    axs[i].hist(thursday_handicap_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(thursday_handicap_categories_labels[i], np.mean(thursday_handicap_categories_hist_scores[i]), np.std(thursday_handicap_categories_hist_scores[i])))
    data = np.array(thursday_handicap_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Thursday Handicap (Event 3) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 42 Thursday handicap Class Histogram
plt.figure()
fig, axs = plt.subplots(len(thursday_handicap_yardages_labels), 1, sharex=True)
for i in range(0,len(thursday_handicap_yardages_labels)):
    axs[i].hist(thursday_handicap_yardages_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(thursday_handicap_yardages_labels[i], np.mean(thursday_handicap_yardages_hist_scores[i]), np.std(thursday_handicap_yardages_hist_scores[i])))
    data = np.array(thursday_handicap_yardages_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Thursday Handicap (Event 3) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 43 Thursday doubles Categories
plt.figure()
plt.pie(thursday_doubles_categories_counts, labels=thursday_doubles_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Thursday Doubles (Event 4) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 44 Thursday doubles Classes
plt.figure()
plt.pie(thursday_doubles_classes_counts, labels=thursday_doubles_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Thursday Doubles (Event 4) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 45 Thursday doubles Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(thursday_doubles_categories_labels), 1, sharex=True)
for i in range(0,len(thursday_doubles_categories_labels)):
    axs[i].hist(thursday_doubles_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(thursday_doubles_categories_labels[i], np.mean(thursday_doubles_categories_hist_scores[i]), np.std(thursday_doubles_categories_hist_scores[i])))
    data = np.array(thursday_doubles_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Thursday doubles (Event 4) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 48 Thursday doubles Class Histogram
plt.figure()
fig, axs = plt.subplots(len(thursday_doubles_classes_labels), 1, sharex=True)
for i in range(0,len(thursday_doubles_classes_labels)):
    axs[i].hist(thursday_doubles_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(thursday_doubles_classes_labels[i], np.mean(thursday_doubles_classes_hist_scores[i]), np.std(thursday_doubles_classes_hist_scores[i])))
    data = np.array(thursday_doubles_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Thursday Doubles (Event 4) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 43 friday doubles Categories
plt.figure()
plt.pie(friday_doubles_categories_counts, labels=friday_doubles_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Doubles (Event 5) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 44 friday doubles Classes
plt.figure()
plt.pie(friday_doubles_classes_counts, labels=friday_doubles_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Doubles (Event 5) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 45 friday doubles Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(friday_doubles_categories_labels), 1, sharex=True)
for i in range(0,len(friday_doubles_categories_labels)):
    axs[i].hist(friday_doubles_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(friday_doubles_categories_labels[i], np.mean(friday_doubles_categories_hist_scores[i]), np.std(friday_doubles_categories_hist_scores[i])))
    data = np.array(friday_doubles_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Friday doubles (Event 5) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 48 friday doubles Class Histogram
plt.figure()
fig, axs = plt.subplots(len(friday_doubles_classes_labels), 1, sharex=True)
for i in range(0,len(friday_doubles_classes_labels)):
    axs[i].hist(friday_doubles_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(friday_doubles_classes_labels[i], np.mean(friday_doubles_classes_hist_scores[i]), np.std(friday_doubles_classes_hist_scores[i])))
    data = np.array(friday_doubles_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Friday Doubles (Event 5) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 49 friday singles Categories
plt.figure()
plt.pie(friday_singles_categories_counts, labels=friday_singles_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Singles (Event 6) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 50 friday singles Classes
plt.figure()
plt.pie(friday_singles_classes_counts, labels=friday_singles_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Singles (Event 6) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 51 friday singles Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(friday_singles_categories_labels), 1, sharex=True)
for i in range(0,len(friday_singles_categories_labels)):
    axs[i].hist(friday_singles_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(friday_singles_categories_labels[i], np.mean(friday_singles_categories_hist_scores[i]), np.std(friday_singles_categories_hist_scores[i])))
    data = np.array(friday_singles_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Preliminary Singles (Event 6) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 52 friday singles Class Histogram
plt.figure()
fig, axs = plt.subplots(len(friday_singles_classes_labels), 1, sharex=True)
for i in range(0,len(friday_singles_classes_labels)):
    axs[i].hist(friday_singles_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(friday_singles_classes_labels[i], np.mean(friday_singles_classes_hist_scores[i]), np.std(friday_singles_classes_hist_scores[i])))
    data = np.array(friday_singles_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Preliminary Singles (Event 6) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 53 friday handicap Categories
plt.figure()
plt.pie(friday_handicap_categories_counts, labels=friday_handicap_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Handicap (Event 7) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 54 friday handicap yardages
plt.figure()
plt.pie(friday_handicap_yardages_counts, labels=friday_handicap_yardages_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Handicap (Event 7) Shooters by Yardage".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 55 friday handicap Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(friday_handicap_categories_labels), 1, sharex=True)
for i in range(0,len(friday_handicap_categories_labels)):
    axs[i].hist(friday_handicap_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(friday_handicap_categories_labels[i], np.mean(friday_handicap_categories_hist_scores[i]), np.std(friday_handicap_categories_hist_scores[i])))
    data = np.array(friday_handicap_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Friday Handicap (Event 7) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 56 friday handicap Class Histogram
plt.figure()
fig, axs = plt.subplots(len(friday_handicap_yardages_labels), 1, sharex=True)
for i in range(0,len(friday_handicap_yardages_labels)):
    axs[i].hist(friday_handicap_yardages_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(friday_handicap_yardages_labels[i], np.mean(friday_handicap_yardages_hist_scores[i]), np.std(friday_handicap_yardages_hist_scores[i])))
    data = np.array(friday_handicap_yardages_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Friday Handicap (Event 7) Scores by Yardage".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 57 championship singles Categories
plt.figure()
plt.pie(championship_singles_categories_counts, labels=championship_singles_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Championship Singles (Event 8) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 58 championship singles Classes
plt.figure()
plt.pie(championship_singles_classes_counts, labels=championship_singles_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Championship Singles (Event 8) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 59 championship singles Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(championship_singles_categories_labels), 1, sharex=True)
for i in range(0,len(championship_singles_categories_labels)):
    axs[i].hist(championship_singles_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(championship_singles_categories_labels[i], np.mean(championship_singles_categories_hist_scores[i]), np.std(championship_singles_categories_hist_scores[i])))
    data = np.array(championship_singles_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 200)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Championship Singles (Event 8) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 60 championship singles Class Histogram
plt.figure()
fig, axs = plt.subplots(len(championship_singles_classes_labels), 1, sharex=True)
for i in range(0,len(championship_singles_classes_labels)):
    axs[i].hist(championship_singles_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(championship_singles_classes_labels[i], np.mean(championship_singles_classes_hist_scores[i]), np.std(championship_singles_classes_hist_scores[i])))
    data = np.array(championship_singles_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 200)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Championship Singles (Event 8) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 61 championship doubles Categories
plt.figure()
plt.pie(championship_doubles_categories_counts, labels=championship_doubles_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Championship Doubles (Event 9) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 62 championship doubles Classes
plt.figure()
plt.pie(championship_doubles_classes_counts, labels=championship_doubles_classes_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Championship Doubles (Event 9) Shooters by Class".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 63 championship doubles Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(championship_doubles_categories_labels), 1, sharex=True)
for i in range(0,len(championship_doubles_categories_labels)):
    axs[i].hist(championship_doubles_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(championship_doubles_categories_labels[i], np.mean(championship_doubles_categories_hist_scores[i]), np.std(championship_doubles_categories_hist_scores[i])))
    if championship_doubles_categories_labels[i] !="LD2":
        data = np.array(championship_doubles_categories_hist_scores[i])
        a, loc, scale = stats.skewnorm.fit(data)
        x = np.linspace(0, 100)
        pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
        axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Championship Doubles (Event 9) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 64 championship doubles Class Histogram
plt.figure()
fig, axs = plt.subplots(len(championship_doubles_classes_labels), 1, sharex=True)
for i in range(0,len(championship_doubles_classes_labels)):
    axs[i].hist(championship_doubles_classes_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(championship_doubles_classes_labels[i], np.mean(championship_doubles_classes_hist_scores[i]), np.std(championship_doubles_classes_hist_scores[i])))
    data = np.array(championship_doubles_classes_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Championship Doubles (Event 9) Scores by Class".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 65 championship handicap Categories
plt.figure()
plt.pie(championship_handicap_categories_counts, labels=championship_handicap_categories_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Preliminary Handicap (Event 10) Shooters by Category".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 66 championship handicap yardages
plt.figure()
plt.pie(championship_handicap_yardages_counts, labels=championship_handicap_yardages_labels, autopct="%1.1f%%")
plt.title("Figure {}: {} Championship Handicap (Even 10) Shooters by Yardage".format(figure, years[0]))
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 67 championship handicap Categories Histograms
plt.figure()
fig, axs = plt.subplots(len(championship_handicap_categories_labels), 1, sharex=True)
for i in range(0,len(championship_handicap_categories_labels)):
    axs[i].hist(championship_handicap_categories_hist_scores[i], alpha=0.5, density=True, bins=10, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(championship_handicap_categories_labels[i], np.mean(championship_handicap_categories_hist_scores[i]), np.std(championship_handicap_categories_hist_scores[i])))
    data = np.array(championship_handicap_categories_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Championship Handicap (Event 10) Scores by Category".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 68 championship handicap Class Histogram
plt.figure()
fig, axs = plt.subplots(len(championship_handicap_yardages_labels), 1, sharex=True)
for i in range(0,len(championship_handicap_yardages_labels)):
    axs[i].hist(championship_handicap_yardages_hist_scores[i], alpha=0.5, bins=10, density=True, label="{0}: \u03BC = {1:.1f}, \u03C3 = {2:.1f}".format(championship_handicap_yardages_labels[i], np.mean(championship_handicap_yardages_hist_scores[i]), np.std(championship_handicap_yardages_hist_scores[i])))
    data = np.array(championship_handicap_yardages_hist_scores[i])
    a, loc, scale = stats.skewnorm.fit(data)
    x = np.linspace(0, 100)
    pdf = stats.skewnorm.pdf(x, a, loc=loc, scale=scale)
    axs[i].plot(x, pdf, 'r-', lw=2, label="Fitted Skew Normal PDF")
    axs[i].legend(loc='upper left')
fig.suptitle("Figure {}: {} Championship Handicap (Event 10) Scores by Yardage".format(figure, years[0]))
fig.set_size_inches(8.5,11)
plt.savefig("figure{}.png".format(figure))
figure += 1
plt.close('all')

#Figure 69: Championship Singles (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Singles"], label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Entries in Championship Singles (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 70: Championship Singles (Extended Years) plot with best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Singles"], label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(df_extended_years_nd["Year"], df_extended_years_nd["Singles"], 1)
x = np.array(df_extended_years_nd["Year"])
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(df_extended_years_nd["Year"]), min(df_extended_years_nd["Singles"]), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Entries in Championship Singles (Per ATA Average Book) \nWith Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 71: Championship Handicap (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Handicap"], label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Entries in Championship Handicap (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 72: Championship Handicap (Extended Years) plot with best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Handicap"], label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(df_extended_years_nd["Year"], df_extended_years_nd["Handicap"], 1)
x = np.array(df_extended_years_nd["Year"])
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(df_extended_years_nd["Year"]), min(df_extended_years_nd["Handicap"]), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Entries in Championship Handicap (Per ATA Average Book) \nWith Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 73: Championship Doubles (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Doubles"], label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Entries in Championship Doubles (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 74: Championship Doubles (Extended Years) plot with best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Doubles"], label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(df_extended_years_nd["Year"], df_extended_years_nd["Doubles"], 1)
x = np.array(df_extended_years_nd["Year"])
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(df_extended_years_nd["Year"]), min(df_extended_years_nd["Doubles"]), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Entries in Championship Doubles (Per ATA Average Book) \nWith Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 75: Championship Total (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Total"], label="Total Shooters", marker='o', color=total_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Total Entries in Championship Events (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 76: Championship Total (Extended Years) plot with best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_nd["Total"], label="Total Shooters", marker='o', color=total_color)
m, b = np.polyfit(df_extended_years_nd["Year"], df_extended_years_nd["Total"], 1)
x = np.array(df_extended_years_nd["Year"])
y = m * x + b
plt.plot(x, y, color=best_fit_color)
plt.text(average(df_extended_years_nd["Year"]), min(df_extended_years_nd["Total"]), "y = {0:,.{2}f}x + {1:,.{2}f}".format(m,b, 1))
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Total Entries in Championship Events (Per ATA Average Book) \nWith Best Fit Line".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 77: Adjacent State Singles (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["South Dakota"], label="South Dakota", marker='o', color="lime")
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["Minnesota"], label="Minnesota", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["Montana"], label="Montana", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["Manitoba"], label="Manitoba", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["Saskatchewan"], label="Saskatchewan", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_singles_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Total Entries in Championship Singles (Per ATA Average Book)\nAdjacent States".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 78: Adjacent State handicap (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["South Dakota"], label="South Dakota", marker='o', color="lime")
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["Minnesota"], label="Minnesota", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["Montana"], label="Montana", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["Manitoba"], label="Manitoba", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["Saskatchewan"], label="Saskatchewan", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_handicap_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Total Entries in Championship Handicap (Per ATA Average Book)\nAdjacent States".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 79: Adjacent State doubles (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["South Dakota"], label="South Dakota", marker='o', color="lime")
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["Minnesota"], label="Minnesota", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["Montana"], label="Montana", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["Manitoba"], label="Manitoba", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["Saskatchewan"], label="Saskatchewan", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_doubles_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Total Entries in Championship Doubles (Per ATA Average Book)\nAdjacent States".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 80: Adjacent State total (Extended Years) without best fit lines
plt.figure()
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["South Dakota"], label="South Dakota", marker='o', color="lime")
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["Minnesota"], label="Minnesota", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["Montana"], label="Montana", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["Manitoba"], label="Manitoba", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["Saskatchewan"], label="Saskatchewan", marker='o')
plt.plot(df_extended_years_nd["Year"], df_extended_years_adjacent_total_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(df_extended_years_nd["Year"])
plt.title("Figure {}: Total Entries in Championship Events (Per ATA Average Book)\nAdjacent States".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')


#Figure 81: All Zones Singles (Extended Years) without best fit lines
plt.figure()
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["Central Zone"], label="Central Zone", marker='o', color="lime")
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["Western Zone"], label="Western Zone", marker='o')
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["Southwestern Zone"], label="Southwestern Zone", marker='o')
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["Southern Zone"], label="Southern Zone", marker='o')
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["Eastern Zone"], label="Eastern Zone", marker='o', color="gold")
plt.plot(extended_years_all_singles_pivot_average_years, df_extended_years_all_singles_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(extended_years_all_singles_pivot_average_years)
plt.title("Figure {}: Average Entries in Singles Championship Events \nBy Zone (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 82: All Zones Singles (Extended Years) without best fit lines
plt.figure()
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["Central Zone"], label="Central Zone", marker='o', color="lime")
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["Western Zone"], label="Western Zone", marker='o')
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["Southwestern Zone"], label="Southwestern Zone", marker='o')
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["Southern Zone"], label="Southern Zone", marker='o')
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["Eastern Zone"], label="Eastern Zone", marker='o', color="gold")
plt.plot(extended_years_all_handicap_pivot_average_years, df_extended_years_all_handicap_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(extended_years_all_handicap_pivot_average_years)
plt.title("Figure {}: Average Entries in Handicap Championship Events \nBy Zone (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 83: All Zones doubles (Extended Years) without best fit lines
plt.figure()
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["Central Zone"], label="Central Zone", marker='o', color="lime")
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["Western Zone"], label="Western Zone", marker='o')
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["Southwestern Zone"], label="Southwestern Zone", marker='o')
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["Southern Zone"], label="Southern Zone", marker='o')
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["Eastern Zone"], label="Eastern Zone", marker='o', color="gold")
plt.plot(extended_years_all_doubles_pivot_average_years, df_extended_years_all_doubles_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(extended_years_all_doubles_pivot_average_years)
plt.title("Figure {}: Average Entries in Doubles Championship Events \nBy Zone (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#Figure 84: All Zones total (Extended Years) without best fit lines
plt.figure()
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["North Dakota"], label="North Dakota", marker='o', color=total_color)
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["Central Zone"], label="Central Zone", marker='o', color="lime")
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["Western Zone"], label="Western Zone", marker='o')
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["Southwestern Zone"], label="Southwestern Zone", marker='o')
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["Southern Zone"], label="Southern Zone", marker='o')
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["Eastern Zone"], label="Eastern Zone", marker='o', color="gold")
plt.plot(extended_years_all_total_pivot_average_years, df_extended_years_all_total_pivot["Average"], label="Average", marker='o', color=best_fit_color)
plt.xlabel("Year")
plt.ylabel("# Shooters")
plt.xticks(extended_years_all_total_pivot_average_years)
plt.title("Figure {}: Average Entries in All Championship Events \nBy Zone (Per ATA Average Book)".format(figure))
plt.legend(loc='upper left')
plt.savefig("figure{}.png".format(figure))
figure +=1
plt.close('all')

#document.save("tables.docx")